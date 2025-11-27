"""
Exporter для конвертации Markdown ТЗ в PDF, DOCX, HTML.
"""
import io
import logging
from typing import Optional
from pathlib import Path

logger = logging.getLogger(__name__)


class TSExporter:
    """Экспортер технических заданий в различные форматы."""
    
    def __init__(self):
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Проверяет наличие необходимых библиотек."""
        self.has_pdfkit = False
        self.has_weasyprint = False
        self.has_docx = False
        self.has_markdown = False
        self.has_pypandoc = False
        self.has_reportlab = False
        
        try:
            import pdfkit
            self.has_pdfkit = True
            logger.info("pdfkit доступен")
        except ImportError:
            logger.warning("pdfkit не установлен. Установите: pip install pdfkit")
        
        try:
            import weasyprint
            # Пробуем импортировать основные модули
            from weasyprint import HTML
            self.has_weasyprint = True
            logger.info("weasyprint доступен")
        except (ImportError, OSError) as e:
            logger.warning(f"weasyprint недоступен: {e}. Установите: pip install weasyprint (на Windows может потребоваться GTK+)")
            self.has_weasyprint = False
        
        try:
            import docx
            self.has_docx = True
            logger.info("python-docx доступен")
        except ImportError:
            logger.warning("python-docx не установлен. Установите: pip install python-docx")
        
        try:
            import markdown
            self.has_markdown = True
            logger.info("markdown доступен")
        except ImportError:
            logger.warning("markdown не установлен. Установите: pip install markdown")
        
        try:
            import pypandoc
            self.has_pypandoc = True
            logger.info("pypandoc доступен")
        except ImportError:
            logger.warning("pypandoc не установлен. Установите: pip install pypandoc")
        
        try:
            import reportlab
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate
            self.has_reportlab = True
            logger.info("reportlab доступен (рекомендуется для Windows)")
        except ImportError:
            logger.warning("reportlab не установлен. Установите: pip install reportlab")
            self.has_reportlab = False
        
        try:
            import reportlab
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            self.has_reportlab = True
            logger.info("reportlab доступен")
        except ImportError:
            logger.warning("reportlab не установлен. Установите: pip install reportlab")
            self.has_reportlab = False
    
    def export_to_html(self, markdown_text: str, include_style: bool = True) -> str:
        """
        Конвертирует Markdown в HTML.
        
        Args:
            markdown_text: Текст в формате Markdown
            include_style: Включить CSS стили
        
        Returns:
            str - HTML документ
        """
        if not self.has_markdown:
            # Простая конвертация без библиотеки
            html = self._simple_markdown_to_html(markdown_text)
        else:
            import markdown as md_lib
            html = md_lib.markdown(markdown_text, extensions=['extra', 'tables', 'codehilite'])
        
        if include_style:
            html = self._wrap_with_style(html)
        
        return html
    
    def export_to_pdf(self, markdown_text: str) -> bytes:
        """
        Конвертирует Markdown в PDF.
        
        Args:
            markdown_text: Текст в формате Markdown
        
        Returns:
            bytes - PDF документ
        """
        # Сначала конвертируем в HTML
        html = self.export_to_html(markdown_text, include_style=True)
        
        # Пробуем разные методы конвертации
        # Пробуем reportlab как альтернативу (работает на Windows)
        if self.has_reportlab:
            try:
                # Используем новый модуль с правильной поддержкой кириллицы
                from app.ts_generator.pdf_exporter import export_markdown_to_pdf
                return export_markdown_to_pdf(markdown_text)
            except ImportError:
                # Fallback на старый метод
                try:
                    return self._export_to_pdf_reportlab(markdown_text)
                except Exception as e:
                    logger.warning(f"Ошибка при использовании старого метода: {e}")
            except Exception as e:
                logger.warning(f"Ошибка при использовании нового экспортера: {e}")
                # Fallback на старый метод
                try:
                    return self._export_to_pdf_reportlab(markdown_text)
                except Exception as e2:
                    logger.warning(f"Ошибка при использовании старого метода: {e2}")
        
        if self.has_pypandoc:
            try:
                import pypandoc
                pdf_bytes = pypandoc.convert_text(
                    markdown_text,
                    'pdf',
                    format='markdown',
                    extra_args=['--pdf-engine=xelatex', '--variable=mainfont:DejaVu Sans']
                )
                if isinstance(pdf_bytes, str):
                    pdf_bytes = pdf_bytes.encode('utf-8')
                return pdf_bytes
            except Exception as e:
                logger.warning(f"Ошибка при использовании pypandoc: {e}")
        
        if self.has_weasyprint:
            try:
                from weasyprint import HTML
                pdf_bytes = HTML(string=html).write_pdf()
                return pdf_bytes
            except Exception as e:
                logger.warning(f"Ошибка при использовании weasyprint: {e}")
        
        if self.has_pdfkit:
            try:
                import pdfkit
                pdf_bytes = pdfkit.from_string(html, False, options={
                    'page-size': 'A4',
                    'margin-top': '0.75in',
                    'margin-right': '0.75in',
                    'margin-bottom': '0.75in',
                    'margin-left': '0.75in',
                    'encoding': "UTF-8",
                    'no-outline': None
                })
                return pdf_bytes
            except Exception as e:
                logger.warning(f"Ошибка при использовании pdfkit: {e}")
        
        raise RuntimeError(
            "Не удалось экспортировать в PDF. Установите одну из библиотек: "
            "reportlab (рекомендуется для Windows), weasyprint, pdfkit или pypandoc"
        )
    
    def export_to_docx(self, markdown_text: str) -> bytes:
        """
        Конвертирует Markdown в DOCX.
        
        Args:
            markdown_text: Текст в формате Markdown
        
        Returns:
            bytes - DOCX документ
        """
        if self.has_pypandoc:
            try:
                import pypandoc
                docx_bytes = pypandoc.convert_text(
                    markdown_text,
                    'docx',
                    format='markdown'
                )
                if isinstance(docx_bytes, str):
                    docx_bytes = docx_bytes.encode('utf-8')
                return docx_bytes
            except Exception as e:
                logger.warning(f"Ошибка при использовании pypandoc: {e}")
        
        if self.has_docx:
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                # Конвертируем Markdown в HTML, затем парсим
                html = self.export_to_html(markdown_text, include_style=False)
                
                # Создаем документ
                doc = Document()
                
                # Парсим HTML и добавляем в документ
                self._html_to_docx(html, doc)
                
                # Сохраняем в bytes
                buffer = io.BytesIO()
                doc.save(buffer)
                return buffer.getvalue()
            except Exception as e:
                logger.warning(f"Ошибка при использовании python-docx: {e}")
        
        raise RuntimeError(
            "Не удалось экспортировать в DOCX. Установите одну из библиотек: "
            "python-docx или pypandoc"
        )
    
    def _simple_markdown_to_html(self, markdown_text: str) -> str:
        """Простая конвертация Markdown в HTML без библиотек."""
        html = markdown_text
        
        # Заголовки
        for i in range(6, 0, -1):
            html = html.replace('#' * i + ' ', f'<h{i}>')
            html = html.replace('\n' + '#' * i + ' ', f'\n<h{i}>')
            # Закрывающие теги
            lines = html.split('\n')
            new_lines = []
            for line in lines:
                if line.startswith(f'<h{i}>') and not line.endswith(f'</h{i}>'):
                    if '\n' in line:
                        parts = line.split('\n', 1)
                        new_lines.append(parts[0] + f'</h{i}>')
                        if len(parts) > 1:
                            new_lines.append(parts[1])
                    else:
                        new_lines.append(line + f'</h{i}>')
                else:
                    new_lines.append(line)
            html = '\n'.join(new_lines)
        
        # Жирный текст
        html = html.replace('**', '<strong>', 1)
        html = html.replace('**', '</strong>', 1)
        while '**' in html:
            html = html.replace('**', '<strong>', 1)
            html = html.replace('**', '</strong>', 1)
        
        # Списки
        lines = html.split('\n')
        new_lines = []
        in_list = False
        for line in lines:
            if line.strip().startswith('- '):
                if not in_list:
                    new_lines.append('<ul>')
                    in_list = True
                new_lines.append(f'<li>{line.strip()[2:]}</li>')
            elif line.strip().startswith(('1. ', '2. ', '3. ', '4. ', '5. ')):
                if not in_list:
                    new_lines.append('<ol>')
                    in_list = True
                new_lines.append(f'<li>{line.strip()[3:]}</li>')
            else:
                if in_list:
                    new_lines.append('</ul>' if '- ' in '\n'.join(new_lines[-10:]) else '</ol>')
                    in_list = False
                new_lines.append(line)
        if in_list:
            new_lines.append('</ul>')
        html = '\n'.join(new_lines)
        
        # Параграфы
        html = html.replace('\n\n', '</p><p>')
        html = '<p>' + html + '</p>'
        
        return html
    
    def _wrap_with_style(self, html: str) -> str:
        """Оборачивает HTML в полный документ со стилями."""
        style = """
        <style>
            body {
                font-family: 'DejaVu Sans', Arial, sans-serif;
                line-height: 1.6;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
                color: #333;
            }
            h1 {
                color: #2c3e50;
                border-bottom: 3px solid #3498db;
                padding-bottom: 10px;
            }
            h2 {
                color: #34495e;
                margin-top: 30px;
                border-bottom: 2px solid #ecf0f1;
                padding-bottom: 5px;
            }
            h3 {
                color: #7f8c8d;
                margin-top: 20px;
            }
            table {
                border-collapse: collapse;
                width: 100%;
                margin: 20px 0;
            }
            table th, table td {
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }
            table th {
                background-color: #3498db;
                color: white;
            }
            code {
                background-color: #f4f4f4;
                padding: 2px 4px;
                border-radius: 3px;
                font-family: 'Courier New', monospace;
            }
            pre {
                background-color: #f4f4f4;
                padding: 10px;
                border-radius: 5px;
                overflow-x: auto;
            }
            ul, ol {
                margin: 10px 0;
                padding-left: 30px;
            }
            li {
                margin: 5px 0;
            }
        </style>
        """
        
        return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Техническое задание</title>
    {style}
</head>
<body>
{html}
</body>
</html>"""
    
    def _html_to_docx(self, html: str, doc: 'Document'):
        """Конвертирует HTML в DOCX документ."""
        from bs4 import BeautifulSoup
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        soup = BeautifulSoup(html, 'html.parser')
        
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'li', 'table']):
            if element.name.startswith('h'):
                level = int(element.name[1])
                heading = doc.add_heading(element.get_text(), level=level)
            elif element.name == 'p':
                p = doc.add_paragraph(element.get_text())
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text(), style='List Bullet')
            elif element.name == 'ol':
                for li in element.find_all('li', recursive=False):
                    doc.add_paragraph(li.get_text(), style='List Number')
            elif element.name == 'table':
                table = doc.add_table(rows=1, cols=len(element.find_all('th')))
                table.style = 'Light Grid Accent 1'
                
                # Заголовки
                headers = element.find_all('th')
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = header.get_text()
                
                # Данные
                for row in element.find_all('tr')[1:]:
                    cells = row.find_all('td')
                    if cells:
                        row_cells = table.add_row().cells
                        for i, cell in enumerate(cells):
                            if i < len(row_cells):
                                row_cells[i].text = cell.get_text()
    
    def _export_to_pdf_reportlab(self, markdown_text: str) -> bytes:
        """Экспорт в PDF через reportlab (работает на Windows)."""
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.fonts import addMapping
        import re
        from io import BytesIO
        import os
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=18)
        
        # Регистрируем шрифты с поддержкой кириллицы
        # Используем системные шрифты Windows
        main_font = 'Helvetica'  # По умолчанию
        bold_font = 'Helvetica-Bold'
        
        # Пробуем загрузить Arial (есть на всех Windows и поддерживает кириллицу)
        try:
            arial_path = r'C:\Windows\Fonts\arial.ttf'
            arial_bold_path = r'C:\Windows\Fonts\arialbd.ttf'
            
            if os.path.exists(arial_path):
                # Регистрируем шрифты
                pdfmetrics.registerFont(TTFont('ArialRU', arial_path))
                main_font = 'ArialRU'
                if os.path.exists(arial_bold_path):
                    pdfmetrics.registerFont(TTFont('ArialRUBold', arial_bold_path))
                    bold_font = 'ArialRUBold'
                else:
                    bold_font = 'ArialRU'
                logger.info("Загружены шрифты Arial для кириллицы")
            else:
                # Пробуем Calibri
                calibri_path = r'C:\Windows\Fonts\calibri.ttf'
                if os.path.exists(calibri_path):
                    pdfmetrics.registerFont(TTFont('CalibriRU', calibri_path))
                    main_font = 'CalibriRU'
                    bold_font = 'CalibriRU'
                    logger.info("Загружены шрифты Calibri для кириллицы")
        except Exception as e:
            logger.warning(f"Не удалось загрузить кириллические шрифты: {e}")
            logger.warning("Используются встроенные шрифты (могут отображать квадраты для кириллицы)")
        
        styles = getSampleStyleSheet()
        story = []
        
        # Создаем кастомные стили с кириллическими шрифтами
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=main_font,
            fontSize=18,
            textColor='#2c3e50',
            spaceAfter=12,
            alignment=TA_LEFT
        )
        
        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontName=main_font,
            fontSize=14,
            textColor='#34495e',
            spaceAfter=10,
            spaceBefore=12
        )
        
        heading3_style = ParagraphStyle(
            'CustomHeading3',
            parent=styles['Heading3'],
            fontName=main_font,
            fontSize=12,
            textColor='#7f8c8d',
            spaceAfter=8,
            spaceBefore=10
        )
        
        # Стиль для обычного текста
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=main_font,
            fontSize=10,
            leading=14
        )
        
        # Убеждаемся, что все стили используют правильный шрифт
        # (на случай, если fontName не применился)
        for style in [title_style, heading2_style, heading3_style, normal_style]:
            style.fontName = main_font
        
        # Парсим Markdown
        lines = markdown_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue
            
            # Заголовки
            if line.startswith('# '):
                text = self._escape_xml(line[2:].strip())
                story.append(Paragraph(text, title_style))
                story.append(Spacer(1, 12))
            elif line.startswith('## '):
                text = self._escape_xml(line[3:].strip())
                story.append(Paragraph(text, heading2_style))
                story.append(Spacer(1, 10))
            elif line.startswith('### '):
                text = self._escape_xml(line[4:].strip())
                story.append(Paragraph(text, heading3_style))
                story.append(Spacer(1, 8))
            elif line.startswith('#### '):
                text = self._escape_xml(line[5:].strip())
                # Создаем стиль для h4
                h4_style = ParagraphStyle(
                    'CustomHeading4',
                    parent=styles['Heading4'],
                    fontName=main_font,
                    fontSize=11
                )
                story.append(Paragraph(text, h4_style))
                story.append(Spacer(1, 6))
            # Жирный текст
            elif '**' in line:
                # Заменяем **текст** на <b>текст</b> правильно
                text = line
                # Находим все пары **
                import re
                # Заменяем **текст** на <b>текст</b>
                text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', text)
                # Убираем оставшиеся **
                text = text.replace('**', '')
                text = self._escape_xml(text)
                story.append(Paragraph(text, normal_style))
                story.append(Spacer(1, 6))
            # Списки
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                # Экранируем для XML
                text = self._escape_xml(text)
                story.append(Paragraph(f"• {text}", normal_style))
                story.append(Spacer(1, 4))
            elif re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line)
                text = self._escape_xml(text)
                story.append(Paragraph(text, normal_style))
                story.append(Spacer(1, 4))
            # Обычный текст
            else:
                # Экранируем для XML
                text = self._escape_xml(line)
                story.append(Paragraph(text, normal_style))
                story.append(Spacer(1, 6))
        
        # Собираем PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def _escape_xml(self, text: str) -> str:
        """Экранирует XML символы для reportlab."""
        return (text
                .replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;')
                .replace("'", '&apos;'))
    
    def _simple_html_to_docx(self, html: str, doc: 'Document'):
        """Простая конвертация HTML в DOCX без BeautifulSoup."""
        from docx.shared import Pt
        
        # Простой парсинг - разбиваем по тегам
        lines = html.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Заголовки
            if line.startswith('<h1>'):
                doc.add_heading(line.replace('<h1>', '').replace('</h1>', ''), level=1)
            elif line.startswith('<h2>'):
                doc.add_heading(line.replace('<h2>', '').replace('</h2>', ''), level=2)
            elif line.startswith('<h3>'):
                doc.add_heading(line.replace('<h3>', '').replace('</h3>', ''), level=3)
            elif line.startswith('<h4>'):
                doc.add_heading(line.replace('<h4>', '').replace('</h4>', ''), level=4)
            elif line.startswith('<h5>'):
                doc.add_heading(line.replace('<h5>', '').replace('</h5>', ''), level=5)
            elif line.startswith('<h6>'):
                doc.add_heading(line.replace('<h6>', '').replace('</h6>', ''), level=6)
            # Параграфы
            elif line.startswith('<p>'):
                text = line.replace('<p>', '').replace('</p>', '')
                # Убираем HTML теги
                text = text.replace('<strong>', '').replace('</strong>', '')
                text = text.replace('<em>', '').replace('</em>', '')
                if text.strip():
                    doc.add_paragraph(text)
            # Списки
            elif line.startswith('<li>'):
                text = line.replace('<li>', '').replace('</li>', '')
                text = text.replace('<strong>', '').replace('</strong>', '')
                if text.strip():
                    doc.add_paragraph(text, style='List Bullet')

